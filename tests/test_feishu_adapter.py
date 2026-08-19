"""FeishuCliAdapter 单元测试（mock subprocess，不依赖真实 lark-cli / 网络）。

覆盖：命令拼接、认证判定、错误分类、run() 对 ok:false 的捕获、
各类能力方法的参数构造、download_media 的 cwd+产物校验。
"""
from __future__ import annotations

import json
import os
import shutil
import subprocess
import sys
import unittest
from unittest.mock import patch

# 允许以脚本方式直接运行（python tests/test_feishu_adapter.py）
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from src.im_adapter.base import BaseIMAdapter  # noqa: E402
from src.im_adapter.base_adapter import BaseIMAdapter as BaseIMAdapterCore  # noqa: E402
from src.im_adapter.errors import (  # noqa: E402
    IMAdapterError,
    IMAdapterRetryableError,
)
from src.im_adapter.feishu import FeishuCliAdapter  # noqa: E402


def _fake(stdout="{}", returncode=0, stderr=""):
    return subprocess.CompletedProcess(args=[], returncode=returncode,
                                       stdout=stdout, stderr=stderr)


class _RunCapture:
    """记录 subprocess.run 的调用参数，并按配置返回假响应。"""

    def __init__(self, response=None, side_effect=None):
        self.calls = []
        self._response = response
        self._side_effect = side_effect

    def __call__(self, *args, **kwargs):
        self.calls.append((args, kwargs))
        if self._side_effect is not None:
            return self._side_effect(*args, **kwargs)
        return self._response

    @property
    def last_cmd(self):
        return self.calls[-1][0][0] if self.calls else None


WHOAMI_READY = json.dumps({
    "profile": "cli_test", "appId": "cli_test", "brand": "feishu",
    "available": True, "tokenStatus": "ready",
    "onBehalfOf": {"userName": "OWNER", "openId": "ou_xxx"},
})


class TestBuildCommand(unittest.TestCase):
    def test_default_cli_is_lark_cli(self):
        a = FeishuCliAdapter()
        self.assertEqual(a.cli_path, "lark-cli")

    def test_inherits_base_and_skeleton(self):
        self.assertTrue(issubclass(FeishuCliAdapter, BaseIMAdapterCore))
        self.assertTrue(issubclass(FeishuCliAdapter, BaseIMAdapter))

    def test_build_no_dry_run(self):
        a = FeishuCliAdapter(dry_run=False)
        cmd = a._build_command(["im", "+messages-send", "--chat-id", "X", "--text", "hi"])
        self.assertEqual(cmd[0], "lark-cli")
        self.assertIn("im", cmd)
        self.assertIn("--chat-id", cmd)
        self.assertNotIn("--dry-run", cmd)

    def test_build_with_dry_run_appends_tail(self):
        a = FeishuCliAdapter(dry_run=True)
        cmd = a._build_command(["im", "+messages-send"])
        self.assertIn("--dry-run", cmd)

    def test_build_with_profile(self):
        a = FeishuCliAdapter(dry_run=False, profile="myprof")
        cmd = a._build_command(["whoami"])
        self.assertIn("--profile", cmd)
        self.assertIn("myprof", cmd)


class TestAuth(unittest.TestCase):
    def test_is_authenticated_true(self):
        cap = _RunCapture(_fake(WHOAMI_READY))
        with patch("subprocess.run", cap):
            self.assertTrue(FeishuCliAdapter(dry_run=False).is_authenticated())

    def test_is_authenticated_false_when_unavailable(self):
        whoami = json.dumps({"available": False, "tokenStatus": "expired"})
        cap = _RunCapture(_fake(whoami))
        with patch("subprocess.run", cap):
            self.assertFalse(FeishuCliAdapter(dry_run=False).is_authenticated())

    def test_is_authenticated_false_on_error(self):
        def boom(*a, **k):
            raise subprocess.CalledProcessError(1, "lark-cli")
        with patch("subprocess.run", boom):
            self.assertFalse(FeishuCliAdapter(dry_run=False).is_authenticated())

    def test_auth_status_normalizes(self):
        cap = _RunCapture(_fake(WHOAMI_READY))
        with patch("subprocess.run", cap):
            st = FeishuCliAdapter(dry_run=False).auth_status()
        self.assertTrue(st["authenticated"])
        self.assertEqual(st["brand"], "feishu")


class TestClassifyError(unittest.TestCase):
    def test_permission_code(self):
        a = FeishuCliAdapter(dry_run=False)
        err = json.dumps({"ok": False, "error": {"code": 99991663, "message": "no perm"}})
        self.assertIs(a._classify_error(err), a._permission_error_class())

    def test_retryable_code(self):
        a = FeishuCliAdapter(dry_run=False)
        err = json.dumps({"ok": False, "error": {"code": 99991400, "message": "rate limit"}})
        self.assertIs(a._classify_error(err), a._retryable_error_class())

    def test_unknown_code_base(self):
        a = FeishuCliAdapter(dry_run=False)
        err = json.dumps({"ok": False, "error": {"code": 99992356, "message": "not exists"}})
        self.assertIs(a._classify_error(err), a._base_error_class())

    def test_keyword_hint_permission(self):
        a = FeishuCliAdapter(dry_run=False)
        self.assertIs(a._classify_error("missing required scope(s): im:x"),
                      a._permission_error_class())

    def test_cross_tenant_app_out_of_chat_classified_permission(self):
        """230002/232010/99992361 是永久权限边界错误，须归类为权限错误（不刷 ERROR、不重试）。"""
        a = FeishuCliAdapter(dry_run=False)
        cases = [
            (230002, "Bot/User can NOT be out of the chat."),
            (232010, "Operator and chat can NOT be in different tenants."),
            (99992361, "open_id cross app"),
        ]
        for code, msg in cases:
            err = json.dumps({"ok": False, "error": {"code": code, "message": msg}})
            self.assertIs(a._classify_error(err), a._permission_error_class(),
                           f"code={code} 应归类为权限错误")

    def test_cross_tenant_message_only_hint_classified_permission(self):
        """生产路径：run() 仅把 message 文本传给 _classify_error（code 解析为 None），须靠 hint 命中。"""
        a = FeishuCliAdapter(dry_run=False)
        self.assertIs(a._classify_error("Operator and chat can NOT be in different tenants."),
                      a._permission_error_class())
        self.assertIs(a._classify_error("open_id cross app"),
                      a._permission_error_class())
        self.assertIs(a._classify_error("Bot/User can NOT be out of the chat."),
                      a._permission_error_class())


class TestRunOkFalse(unittest.TestCase):
    def test_ok_false_with_exit_0_raises(self):
        # lark-cli 可能退出码 0 但 {"ok":false}
        resp = json.dumps({"ok": False, "error": {"code": 99991400, "message": "rate"}})
        cap = _RunCapture(_fake(resp, returncode=0))
        with patch("subprocess.run", cap):
            with self.assertRaises(IMAdapterRetryableError):
                FeishuCliAdapter(dry_run=False).run(["im", "+x"])

    def test_ok_true_returns_data(self):
        resp = json.dumps({"ok": True, "data": {"x": 1}})
        cap = _RunCapture(_fake(resp))
        with patch("subprocess.run", cap):
            d = FeishuCliAdapter(dry_run=False).run(["whoami"])
        self.assertEqual(d["data"]["x"], 1)


class TestCapabilityArgs(unittest.TestCase):
    def _adapter(self, response):
        cap = _RunCapture(_fake(response))
        patcher = patch("subprocess.run", cap)
        patcher.start()
        self.addCleanup(patcher.stop)
        return FeishuCliAdapter(dry_run=False), cap

    def test_chat_message_send_text(self):
        a, cap = self._adapter(json.dumps({"ok": True, "data": {"message_id": "om_1"}}))
        r = a.chat_message_send(group="oc_g", text="hi", uuid="u1")
        cmd = cap.last_cmd
        self.assertIn("--chat-id", cmd)
        self.assertIn("oc_g", cmd)
        self.assertIn("--text", cmd)
        self.assertIn("hi", cmd)
        self.assertIn("--idempotency-key", cmd)
        self.assertIn("u1", cmd)
        self.assertEqual(r["data"]["message_id"], "om_1")

    def test_chat_message_send_markdown_user(self):
        a, cap = self._adapter(json.dumps({"ok": True, "data": {}}))
        a.chat_message_send(user="ou_x", title="t", text="b", msg_type="markdown")
        cmd = cap.last_cmd
        self.assertIn("--user-id", cmd)
        self.assertIn("ou_x", cmd)
        self.assertIn("--markdown", cmd)
        self.assertIn("b", cmd)

    def test_chat_message_send_at_all_group(self):
        a, cap = self._adapter(json.dumps({"ok": True, "data": {}}))
        a.chat_message_send(group="oc_g", text="hi", at_all=True)
        cmd = cap.last_cmd
        self.assertIn("--text", cmd)
        joined = " ".join(cmd)
        self.assertIn("hi", joined)
        # 历史 bug：at_all 参数被静默丢弃；修复后应内联 <at> 语法
        self.assertIn('<at user_id="all"></at>', joined)

    def test_chat_message_send_at_open_ids_group(self):
        a, cap = self._adapter(json.dumps({"ok": True, "data": {}}))
        a.chat_message_send(group="oc_g", text="hi",
                            at_open_dingtalk_ids="ou_a,ou_b")
        joined = " ".join(cap.last_cmd)
        self.assertIn('<at user_id="ou_a"></at>', joined)
        self.assertIn('<at user_id="ou_b"></at>', joined)

    def test_chat_message_send_at_ignored_for_user(self):
        # 契约：@ 仅群聊生效；发给单聊用户时不应追加 <at>
        a, cap = self._adapter(json.dumps({"ok": True, "data": {}}))
        a.chat_message_send(user="ou_x", text="hi", at_all=True)
        cmd = cap.last_cmd
        self.assertNotIn("<at", cmd)

    def test_chat_message_reply_text_thread_uuid(self):
        a, cap = self._adapter(json.dumps({"ok": True, "data": {"message_id": "om_9"}}))
        r = a.chat_message_reply(message_id="om_9", text="回复内容",
                                 reply_in_thread=True, uuid="id-1")
        joined = " ".join(cap.last_cmd)
        self.assertIn("im", cap.last_cmd)
        self.assertIn("+messages-reply", cap.last_cmd)
        self.assertIn("--message-id", cap.last_cmd)
        self.assertIn("om_9", cap.last_cmd)
        self.assertIn("--reply-in-thread", cap.last_cmd)
        self.assertIn("--text", cap.last_cmd)
        self.assertIn("回复内容", joined)
        self.assertIn("--idempotency-key", cap.last_cmd)
        self.assertIn("id-1", cap.last_cmd)
        self.assertEqual(r["data"]["message_id"], "om_9")

    def test_chat_message_reply_markdown_no_thread(self):
        a, cap = self._adapter(json.dumps({"ok": True, "data": {}}))
        a.chat_message_reply(message_id="om_x", text="b", msg_type="markdown")
        joined = " ".join(cap.last_cmd)
        self.assertIn("--markdown", cap.last_cmd)
        self.assertIn("b", joined)
        self.assertNotIn("--reply-in-thread", cap.last_cmd)  # 默认不进话题

    def test_chat_message_reply_requires_message_id(self):
        a, _ = self._adapter(json.dumps({"ok": True, "data": {}}))
        with self.assertRaises(ValueError):
            a.chat_message_reply(text="hi")

    def test_chat_message_list_group(self):
        a, cap = self._adapter(json.dumps({"ok": True, "data": {"messages": []}}))
        a.chat_message_list("oc_g", "2024-01-01", limit=10)
        cmd = cap.last_cmd
        self.assertIn("im", cmd)
        self.assertIn("+chat-messages-list", cmd)
        self.assertIn("--chat-id", cmd)
        self.assertIn("oc_g", cmd)
        self.assertIn("--start", cmd)
        self.assertIn("2024-01-01", cmd)
        self.assertIn("--order", cmd)
        self.assertIn("asc", cmd)
        self.assertIn("--page-size", cmd)
        self.assertIn("10", cmd)

    def test_chat_message_list_without_cached_result(self):
        """chat_message_list 不接受 cached_result；参数移除后调用正常。"""
        a, cap = self._adapter(json.dumps({"ok": True, "data": {"messages": [
            {"message_id": "m1"}]}}))
        msgs = a.chat_message_list("oc_g", "2024-01-01")
        self.assertEqual(len(msgs), 1)
        self.assertEqual(msgs[0]["message_id"], "m1")

    def test_chat_message_list_direct(self):
        a, cap = self._adapter(json.dumps({"ok": True, "data": {"messages": []}}))
        a.chat_message_list_direct(user_id="ou_x")
        cmd = cap.last_cmd
        self.assertIn("--user-id", cmd)
        self.assertIn("ou_x", cmd)

    def test_unread_conversations_extracts_chats(self):
        resp = json.dumps({"ok": True, "data": {"chats": [
            {"chat_id": "oc_1", "name": "g", "chat_mode": "group"}]}})
        a, cap = self._adapter(resp)
        chats = a.chat_message_list_unread_conversations(5)
        self.assertEqual(len(chats), 1)
        self.assertEqual(chats[0]["chat_id"], "oc_1")

    def test_contact_get_self(self):
        a, cap = self._adapter(json.dumps({"ok": True, "data": {"user": {"name": "x"}}}))
        a.contact_user_get_self()
        cmd = cap.last_cmd
        self.assertIn("contact", cmd)
        self.assertIn("+get-user", cmd)

    def test_contact_search(self):
        a, cap = self._adapter(json.dumps({"ok": True, "data": {"items": []}}))
        a.contact_user_search("张三")
        cmd = cap.last_cmd
        self.assertIn("contact", cmd)
        self.assertIn("+search-user", cmd)
        self.assertIn("--query", cmd)
        self.assertIn("张三", cmd)
        self.assertIn("--as", cmd)
        self.assertIn("user", cmd)

    def test_mark_read_uses_api_patch(self):
        a, cap = self._adapter(json.dumps({"ok": True, "data": {}}))
        a.mark_read("oc_x", "om_y")
        cmd = cap.last_cmd
        self.assertIn("api", cmd)
        self.assertIn("PATCH", cmd)
        self.assertIn("/open-apis/im/v1/chats/oc_x/readed", cmd)
        # --data 含 message_id
        data_idx = cmd.index("--data")
        self.assertIn("om_y", cmd[data_idx + 1])


class TestChatMessageListAllWhitelist(unittest.TestCase):
    """白名单模式：跳过 +chat-list 全量翻页，仅拉已知相关会话 + 单次最近活跃探嗅。"""

    def _adapter_with_side_effect(self, side_effect):
        cap = _RunCapture(side_effect=side_effect)
        patcher = patch("subprocess.run", cap)
        patcher.start()
        self.addCleanup(patcher.stop)
        return FeishuCliAdapter(dry_run=False), cap

    @staticmethod
    def _chat_id_of(cmd):
        if "--chat-id" in cmd:
            return cmd[cmd.index("--chat-id") + 1]
        return None

    def test_whitelist_mode_skips_full_pagination(self):
        def se(*args, **kwargs):
            cmd = args[0]
            if "+chat-list" in cmd:
                # 单次最近活跃探嗅：返回 1 个最近群（不翻页，无 page_token）
                return _fake(json.dumps({"ok": True, "data": {"chats": [
                    {"chat_id": "oc_peek", "name": "PeekGroup", "chat_mode": "group"}
                ]}}))
            if "+chat-messages-list" in cmd:
                return _fake(json.dumps({"ok": True, "data": {"messages": [
                    {"message_id": "om_1", "sender_id": "ou_x", "content": "hi",
                     "msg_type": "text", "create_time": "2024-01-01 12:00:00"}
                ]}}))
            return _fake("{}")

        a, cap = self._adapter_with_side_effect(se)
        result = a.chat_message_list_all(
            "2024-01-01 00:00:00", "2024-01-02 00:00:00",
            chat_ids=["oc_a"],
            chat_meta={"oc_a": {"title": "A", "chat_mode": "p2p", "singleChat": True}},
        )

        # 关键：白名单模式只调一次 +chat-list（探嗅），不进入全量翻页循环
        chat_list_calls = [c for c, k in cap.calls if "+chat-list" in c[0]]
        self.assertEqual(len(chat_list_calls), 1,
                         "白名单模式不应全量翻页 +chat-list")

        # 对白名单会话 oc_a 与探嗅到的 oc_peek 都拉了消息
        pulled = {self._chat_id_of(c[0]) for c, k in cap.calls
                  if "+chat-messages-list" in c[0]}
        self.assertIn("oc_a", pulled)
        self.assertIn("oc_peek", pulled)

        # 组装结果正确：两个会话都在，singleChat 来自 chat_meta
        conv_list = result.get("conversationMessagesList", [])
        by_id = {c["openConversationId"]: c for c in conv_list}
        self.assertIn("oc_a", by_id)
        self.assertIn("oc_peek", by_id)
        self.assertTrue(by_id["oc_a"]["singleChat"],
                        "白名单 singleChat 应来自 chat_meta")
        self.assertFalse(by_id["oc_peek"]["singleChat"],
                         "探嗅到的群聊应判定为 group")

    def test_full_mode_still_paginates(self):
        state = {"n": 0}

        def se(*args, **kwargs):
            cmd = args[0]
            if "+chat-list" in cmd:
                state["n"] += 1
                if state["n"] == 1:
                    return _fake(json.dumps({"ok": True, "data": {
                        "chats": [{"chat_id": "oc_1", "name": "g",
                                   "chat_mode": "group"}],
                        "page_token": "TOK"}}))
                return _fake(json.dumps({"ok": True, "data": {"chats": []}}))
            if "+chat-messages-list" in cmd:
                return _fake(json.dumps({"ok": True, "data": {"messages": []}}))
            return _fake("{}")

        a, cap = self._adapter_with_side_effect(se)
        # 不传 chat_ids → 全量模式，应触发翻页
        a.chat_message_list_all("2024-01-01 00:00:00", "2024-01-02 00:00:00")
        chat_list_calls = [c for c, k in cap.calls if "+chat-list" in c[0]]
        self.assertEqual(len(chat_list_calls), 2,
                         "全量模式（无 chat_ids）应保持 +chat-list 翻页行为")

    def test_permission_error_collected_as_blocked_chats(self):
        """遍历消息时命中永久权限错误（跨租户/跨app/已退群）的会话，应被收集进
        blocked_chats 返回，交由轮询器拉黑——不应抛异常、不应进入消息列表。
        """
        def se(*args, **kwargs):
            cmd = args[0]
            if "+chat-messages-list" in cmd and "--chat-id" in cmd:
                cid = cmd[cmd.index("--chat-id") + 1]
                if cid == "oc_dead":
                    # 模拟 232010 跨租户：lark-cli exit 1 + stderr 含 hint，
                    # run() 归类为 IMAdapterPermissionError 并抛出
                    return _fake(
                        json.dumps({"ok": False, "error": {
                            "code": 232010,
                            "message": "Operator and chat can NOT be in different tenants."}}),
                        returncode=1,
                        stderr="Operator and chat can NOT be in different tenants.")
                return _fake(json.dumps({"ok": True, "data": {"messages": [
                    {"message_id": "om_1", "sender_id": "ou_x", "content": "hi",
                     "msg_type": "text", "create_time": "2024-01-01 12:00:00"}]}}))
            if "+chat-list" in cmd:
                return _fake(json.dumps({"ok": True, "data": {"chats": []}}))
            return _fake("{}")

        a, cap = self._adapter_with_side_effect(se)
        result = a.chat_message_list_all(
            "2024-01-01 00:00:00", "2024-01-02 00:00:00",
            chat_ids=["oc_dead", "oc_good"],
            chat_meta={
                "oc_dead": {"title": "DeadChat", "chat_mode": "p2p", "singleChat": True},
                "oc_good": {"title": "GoodChat", "chat_mode": "p2p", "singleChat": True},
            },
        )
        blocked = result.get("blocked_chats", [])
        self.assertEqual(len(blocked), 1, "仅 oc_dead 应被收集为待拉黑")
        self.assertEqual(blocked[0]["chat_id"], "oc_dead")
        self.assertEqual(blocked[0]["title"], "DeadChat")
        self.assertEqual(blocked[0]["chat_type"], "single")
        self.assertIn("different tenants", blocked[0]["error"])

        # oc_good 正常消息应进入结果；oc_dead 不应产生消息列表
        conv_list = result.get("conversationMessagesList", [])
        by_id = {c["openConversationId"]: c for c in conv_list}
        self.assertIn("oc_good", by_id)
        self.assertNotIn("oc_dead", by_id, "死会话不应产生消息列表")


class TestDownloadMedia(unittest.TestCase):
    def test_download_writes_and_returns_path(self):
        out_dir = "/tmp/_feishu_dl"
        os.makedirs(out_dir, exist_ok=True)
        out_path = os.path.join(out_dir, "img.png")

        def side_effect(*args, **kwargs):
            # 模拟 lark-cli 把文件写到 cwd 下的 basename
            cwd = kwargs.get("cwd", ".")
            name = os.path.basename(out_path)
            with open(os.path.join(cwd, name), "w") as f:
                f.write("binary")
            return _fake("{}", returncode=0)

        cap = _RunCapture(side_effect=side_effect)
        with patch("subprocess.run", cap):
            res = FeishuCliAdapter(dry_run=False).download_media(
                media_id="img_abc", message_id="om_1",
                conversation_id="oc_x", output_path=out_path)
        self.assertEqual(res, out_path)
        self.assertTrue(os.path.exists(out_path))
        # 验证传了 cwd = 输出目录、--output = 文件名
        kw = cap.calls[-1][1]
        self.assertEqual(kw.get("cwd"), out_dir)
        self.assertIn("--output", cap.last_cmd)
        self.assertIn(os.path.basename(out_path), cap.last_cmd)

    def test_download_missing_file_raises(self):
        def side_effect(*args, **kwargs):
            # 命令成功但文件没生成
            return _fake("{}", returncode=0)
        cap = _RunCapture(side_effect=side_effect)
        with patch("subprocess.run", cap):
            with self.assertRaises(IMAdapterError):
                FeishuCliAdapter(dry_run=False, retries=0).download_media(
                    media_id="img_abc", message_id="om_1",
                    conversation_id="oc_x", output_path="/tmp/_feishu_dl/missing.png")

    def test_download_resource_type_video(self):
        # 历史 bug：非 img_ 一律按 file 处理，导致视频(va_)下载命令 --type 错误；
        # 修复后应按 media_key 前缀推断为 video
        out_path = "/tmp/_feishu_dl/vid.bin"

        def side_effect(*args, **kwargs):
            cwd = kwargs.get("cwd", ".")
            os.makedirs(cwd, exist_ok=True)
            with open(os.path.join(cwd, os.path.basename(out_path)), "w") as f:
                f.write("x")
            return _fake("{}", returncode=0)

        cap = _RunCapture(side_effect=side_effect)
        with patch("subprocess.run", cap):
            FeishuCliAdapter(dry_run=False).download_media(
                media_id="va_abc", message_id="om_1",
                conversation_id="oc_x", output_path=out_path)
        self.assertIn("--type", cap.last_cmd)
        self.assertIn("video", cap.last_cmd)


class TestMediaUpload(unittest.TestCase):
    def test_returns_path_when_exists(self):
        p = "/tmp/_feishu_up.txt"
        with open(p, "w") as f:
            f.write("x")
        self.assertEqual(FeishuCliAdapter(dry_run=False).media_upload(p), p)

    def test_missing_path_raises(self):
        with self.assertRaises(ValueError):
            FeishuCliAdapter(dry_run=False).media_upload("/tmp/nope_123.png")


class TestDocSearch(unittest.TestCase):
    """doc_search 解析：知识库(wiki)里的文件节点需取出底层 Drive 文件 token。"""

    def _search(self, data):
        cap = _RunCapture(_fake(json.dumps(data)))
        with patch("subprocess.run", cap):
            return FeishuCliAdapter(dry_run=False).doc_search("连接配置指南")

    def test_wiki_file_node_unwraps_real_file_token(self):
        resp = {
            "ok": True,
            "data": {
                "results": [{
                    "entity_type": "WIKI",
                    "title_highlighted": "打印机连接配置指南.md",
                    "result_meta": {
                        "token": "R63zwoRhliUUHNkqf1OcVpzunfg",
                        "doc_types": "FILE",
                        "file_type": "md",
                        "icon_info": json.dumps({
                            "type": 0, "key": "", "obj_type": 12,
                            "file_type": "md",
                            "token": "J8KsbWlR9oqUx1xPvVSc3xXwn0e",
                            "version": 1,
                        }),
                        "url": "https://open-claw.feishu.cn/wiki/R63zwoRhliUUHNkqf1OcVpzunfg",
                    },
                }]
            },
        }
        docs = self._search(resp)
        self.assertEqual(len(docs), 1)
        self.assertEqual(docs[0]["doc_token"], "J8KsbWlR9oqUx1xPvVSc3xXwn0e")
        self.assertEqual(docs[0]["token"], "J8KsbWlR9oqUx1xPvVSc3xXwn0e")
        self.assertEqual(docs[0]["entity_type"], "FILE")
        self.assertTrue(docs[0]["importable"])

    def test_wiki_doc_node_keeps_wiki_token(self):
        # 知识库里的普通文档（doc），仍走 wiki 导出，token 保持 wiki 节点 token
        resp = {
            "ok": True,
            "data": {
                "results": [{
                    "entity_type": "WIKI",
                    "title_highlighted": "会议纪要",
                    "result_meta": {
                        "token": "wikiDocTok123",
                        "doc_types": "DOC",
                        "icon_info": json.dumps({"type": 0, "token": "wikiDocTok123"}),
                        "url": "https://open-claw.feishu.cn/wiki/wikiDocTok123",
                    },
                }]
            },
        }
        docs = self._search(resp)
        self.assertEqual(docs[0]["doc_token"], "wikiDocTok123")
        self.assertEqual(docs[0]["entity_type"], "WIKI")

    def test_wiki_file_icon_info_malformed_keeps_wiki_token(self):
        # icon_info 损坏时不崩溃，退回 wiki 节点 token（安全降级，不 404 也不异常）
        resp = {
            "ok": True,
            "data": {
                "results": [{
                    "entity_type": "WIKI",
                    "title_highlighted": "x.md",
                    "result_meta": {
                        "token": "wikiTok",
                        "doc_types": "FILE",
                        "icon_info": "not-json",
                    },
                }]
            },
        }
        docs = self._search(resp)
        self.assertEqual(docs[0]["doc_token"], "wikiTok")
        self.assertEqual(docs[0]["entity_type"], "WIKI")


class TestRunCliFile(unittest.TestCase):
    """_run_cli_file 必须能定位下载分支产出的文件（名字恰好等于 output_base）。"""

    @patch("subprocess.run")
    def test_picks_file_named_like_output_base(self, mock_run):
        import tempfile
        tmp = tempfile.mkdtemp()
        out_name = "feishu_dl"

        def _side(cmd, *a, **k):
            cwd = k.get("cwd")
            p = os.path.join(cwd, out_name)
            open(p, "wb").write(b"hello")
            return subprocess.CompletedProcess(args=cmd, returncode=0, stdout="", stderr="")

        mock_run.side_effect = _side
        a = FeishuCliAdapter(dry_run=False)
        got = a._run_cli_file(
            ["drive", "+download", "--file-token", "X", "--output", "./" + out_name],
            tmp, output_base=out_name)
        self.assertTrue(got.endswith(out_name))
        self.assertTrue(os.path.isfile(got))
        shutil.rmtree(tmp, ignore_errors=True)


class TestDocReadFallback(unittest.TestCase):
    """非 docx 文档导入回退链路（不依赖真实 lark-cli）。"""

    def setUp(self):
        self.a = FeishuCliAdapter()

    def test_detect_rename_zip_to_docx(self):
        import zipfile
        import tempfile
        p = tempfile.mktemp()
        with zipfile.ZipFile(p, "w") as z:
            z.writestr("word/document.xml", "<w/>")
        out = self.a._detect_and_rename(p)
        self.assertTrue(out.endswith(".docx"))
        os.remove(out)

    def test_detect_rename_pdf(self):
        import tempfile
        p = tempfile.mktemp()
        open(p, "wb").write(b"%PDF-1.4 fake")
        out = self.a._detect_and_rename(p)
        self.assertTrue(out.endswith(".pdf"))
        os.remove(out)

    def test_extract_docx(self):
        import docx
        import tempfile
        p = tempfile.mktemp(suffix=".docx")
        d = docx.Document()
        d.add_paragraph("抽取测试文本ABC")
        d.save(p)
        self.assertIn("抽取测试文本ABC", self.a._extract_text_from_file(p))
        os.remove(p)

    def test_extract_text_file_direct(self):
        import tempfile
        p = tempfile.mktemp(suffix=".md")
        open(p, "w", encoding="utf-8").write("# 标题\n正文内容")
        self.assertIn("正文内容", self.a._extract_text_from_file(p))
        os.remove(p)

    @patch.object(FeishuCliAdapter, "_run_cli_file")
    def test_fallback_routes_export_for_sheet(self, mock_run):
        import openpyxl
        import tempfile
        p = tempfile.mktemp(suffix=".xlsx")
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.append(["姓名", "部门"])
        ws.append(["张三", "研发"])
        wb.save(p)
        mock_run.return_value = p
        res = self.a._doc_read_fallback("sheet_tok", entity_type="SHEET")
        self.assertNotIn("error", res)
        self.assertEqual(res["imported_via"], "export")
        self.assertIn("张三", res["content"])
        os.remove(p)

    @patch.object(FeishuCliAdapter, "_run_cli_file")
    def test_fallback_routes_download_for_file(self, mock_run):
        import docx
        import tempfile
        p = tempfile.mktemp(suffix=".docx")
        d = docx.Document()
        d.add_paragraph("文件类型抽取OK")
        d.save(p)
        mock_run.return_value = p
        res = self.a._doc_read_fallback("file_tok", entity_type="FILE")
        self.assertNotIn("error", res)
        self.assertEqual(res["imported_via"], "download")
        self.assertIn("文件类型抽取OK", res["content"])
        os.remove(p)


class TestChatMessageSendFallback(unittest.TestCase):
    """飞书 chat_message_send user→bot 降级覆盖 230027/230002/404 等场景。"""

    def _adapter_with_side_effect(self, side_effect):
        cap = _RunCapture(side_effect=side_effect)
        patcher = patch("subprocess.run", cap)
        patcher.start()
        self.addCleanup(patcher.stop)
        return FeishuCliAdapter(dry_run=False), cap

    def _user_fail_resp(self, msg: str) -> str:
        return json.dumps({
            "ok": False, "identity": "user",
            "error": {"message": msg},
        })

    def test_404_not_found_fallback_to_bot(self):
        """user 身份发群聊 404 时自动降级为 bot 身份重试。"""
        def se(*args, **kwargs):
            cmd = args[0]
            if "+messages-send" in cmd:
                if "--as" not in cmd:
                    # user 身份失败
                    return _fake(self._user_fail_resp("HTTP 404: 404 page not found"), returncode=1)
                # bot 身份成功
                return _fake(json.dumps({"ok": True, "data": {"message_id": "om_bot"}}))
            return _fake("{}")

        a, cap = self._adapter_with_side_effect(se)
        r = a.chat_message_send(group="oc_g", text="hi", uuid="u1")
        self.assertEqual(r["data"]["message_id"], "om_bot")
        # 确认调了两次：user 失败 + bot 重试
        send_calls = [c for c, k in cap.calls if "+messages-send" in c[0]]
        self.assertEqual(len(send_calls), 2)
        self.assertIn("--as", send_calls[1][0])
        self.assertIn("bot", send_calls[1][0])

    def test_230027_fallback_to_bot(self):
        def se(*args, **kwargs):
            cmd = args[0]
            if "+messages-send" in cmd and "--as" not in cmd:
                return _fake(self._user_fail_resp("code 230027"), returncode=1)
            return _fake(json.dumps({"ok": True, "data": {}}))

        a, _ = self._adapter_with_side_effect(se)
        r = a.chat_message_send(group="oc_g", text="hi")
        self.assertTrue(r.get("ok"))

    def test_non_fallback_error_reraised(self):
        """非降级错误（如 500 内部错误）直接抛出，不重试。"""
        def se(*args, **kwargs):
            return _fake(self._user_fail_resp("internal server error 500"), returncode=1)

        a, _ = self._adapter_with_side_effect(se)
        with self.assertRaises(IMAdapterError):
            a.chat_message_send(group="oc_g", text="hi")

    def test_fallback_disabled_no_retry(self):
        """_disable_bot_fallback=True 时禁止降级重试。"""
        def se(*args, **kwargs):
            return _fake(self._user_fail_resp("code 230027"), returncode=1)

        a, _ = self._adapter_with_side_effect(se)
        a._disable_bot_fallback = True
        with self.assertRaises(IMAdapterError):
            a.chat_message_send(group="oc_g", text="hi")


if __name__ == "__main__":
    unittest.main(verbosity=2)
